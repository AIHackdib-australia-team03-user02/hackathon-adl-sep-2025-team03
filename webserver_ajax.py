import os
import shutil
import uuid
import docx
from fastapi import FastAPI, File, UploadFile, Request, BackgroundTasks, Form
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
import openpyxl
import asyncio
import threading
from main import planning_agent, blueprint_search_agent, data_analyst_agent, remediation_agent, model_client, termination, selector_prompt, selector_func
from autogen_agentchat.teams import SelectorGroupChat
from autogen_agentchat.ui import Console
from datetime import datetime
import markdown2

templates = Jinja2Templates(directory="templates")
app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
progress_store = {}
stop_requested = False

@app.get("/", response_class=HTMLResponse)
def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

from fastapi.responses import FileResponse
import tempfile
try:
    from docx import Document
except ImportError:
    Document = None
# Add endpoint to generate and serve DOCX file
@app.get("/download-docx/{task_id}")
def download_docx(task_id: str):
    if Document is None:
        return JSONResponse({"error": "python-docx not installed"}, status_code=500)
    items = progress_store.get(task_id, [])
    doc = Document()
    doc.add_heading("Remediation Results", 0)
    for line in items:
        if line.startswith("__COMPLETE__:"):
            continue
        guideline, summary = line.split(":", 1) if ":" in line else (line, "")
        doc.add_heading(guideline.strip(), level=1)
        doc.add_paragraph(summary.strip())
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    return FileResponse(tmp_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="remediation_results.docx")


@app.post("/upload", response_class=JSONResponse)
async def upload(request: Request, background_tasks: BackgroundTasks = None, file: UploadFile = File(None), use_default: str = Form(None)):
    if use_default == '1':
        file_path = os.path.join("input_spreadsheet", "remediation_annex.xlsx")
    else:
        if file is None:
            return {"error": "No file uploaded and 'use default' not selected."}
        file_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    wb = openpyxl.load_workbook(file_path, data_only=True)
    ws = wb.active
    criteria = []
    for row in ws.iter_rows(min_row=2):
        val = row[14].value
        if val and str(val).strip():
            criteria.append(str(val).strip())
    task_id = str(uuid.uuid4())
    progress_store[task_id] = []
    background_tasks.add_task(run_team_background, criteria, task_id)
    return {"task_id": task_id}

def render_result_item(line):
    guideline, summary = line.split(":", 1) if ":" in line else (line, "")
    summary_clean = summary.strip()
    # Do NOT remove TERMINATE so polling can detect completion
    paragraphs = [p.strip() for p in summary_clean.split("\n") if p.strip()]
    last_two = paragraphs[-2:] if len(paragraphs) >= 2 else paragraphs
    color = None
    for p in last_two:
        if "GREEN" in p:
            color = "green"
        if "RED" in p:
            color = "red"
    li_class = "" if color == "green" else "red"
    summary_html = markdown2.markdown(summary_clean)
    return f'<li class="{li_class}"><span class="guideline">{guideline.strip()}</span><div class="summary">{summary_html}</div></li>'



async def run_team(criteria, task_id):
    global stop_requested
    # Create a fresh team and console for each run
    team = SelectorGroupChat(
        [planning_agent, blueprint_search_agent, data_analyst_agent, remediation_agent],
        model_client=model_client,
        termination_condition=termination,
        selector_prompt=selector_prompt,
        selector_func=selector_func,
        allow_repeated_speaker=True,
        max_turns=15,
    )
    from autogen_agentchat.ui import Console
    for guideline_description in criteria:
        if stop_requested:
            progress_store[task_id].append(f"{guideline_description}: [Stopped by user]")
            break
        task = f"Determine if this criteria has been satisfied with the current setup scripts: '{guideline_description}'"
        await team.reset()
        result = await Console(team.run_stream(task=task))
        if hasattr(result, "messages") and result.messages:
            remediation_html = ""
            if "RED" in result.messages[-1].content:
                remediation_steps = result.messages[-2].content
                remediation_html = f'<span><small class="remediation-steps">{remediation_steps}</small></span><br>'
            summary:str = result.messages[-1].content
            final_message = remediation_html + summary
            progress_store[task_id].append(f"{guideline_description}: {final_message}")
        else:
            progress_store[task_id].append(f"{guideline_description}: [No summary found]")
    # Add a special event at the end to signal completion, but allow polling to continue
    progress_store[task_id].append("__COMPLETE__: All criteria processed.")
    team.reset()

# Wrapper for background task
def run_team_background(criteria, task_id):
    global stop_requested
    stop_requested = False
    threading.Thread(target=lambda: asyncio.run(run_team(criteria, task_id)), daemon=True).start()
@app.post("/stop-tasks", response_class=JSONResponse)
def stop_tasks():
    global stop_requested
    stop_requested = True
    return {"status": "stopping background tasks"}

@app.get("/progress/{task_id}", response_class=JSONResponse)
def get_progress(task_id: str):
    items = progress_store.get(task_id, [])
    html_items = [render_result_item(line) for line in items if not line.startswith("__COMPLETE__:")]
    # Always return items, never signal done
    return {"items": html_items}

if __name__ == "__main__":
    import uvicorn
    import os
    port = int(os.environ.get("PORT", 8000))
    print("Starting server on port", port)
    uvicorn.run(app, host="0.0.0.0", port=port)
